# -*- coding: utf-8 -*-
"""把 5 个批次的论文总结 Markdown 转换为排版美观、通俗易懂的 docx。
结构：封面（R1 配方 / DS-1 配色）→ 目录（罗马页码）→ 正文（阿拉伯页码）。
每篇论文：标题 → 领域/概述信息表 → 通俗解读 → 方法论总结 → 英语词汇表。"""
import re, os, math, sys
from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"D:\AI research\Cell Press_20260901"
SUM = os.path.join(BASE, "_summaries")
OUT = os.path.join(BASE, "docx")
os.makedirs(OUT, exist_ok=True)

# DS-1 Deep Sea 配色
BG = "0B1C2C"; ACCENT = "529286"
C_TITLE = "FFFFFF"; C_SUB = "B0B8C0"; C_META = "90989F"; C_FOOT = "687078"
T_HEADBG = "529286"; T_INNER = "BECFCC"; T_SURF = "E8ECEB"
H1_COLOR = "2F6B62"

# ---------------- 通俗解读（48 篇） ----------------
PLAIN = {
1: "如今电动车电池越做越大，但电极生产还停留在“把材料调成浆糊、涂开、烘干”的湿法工艺，又费电又用有毒溶剂。这项研究改用“干粉直接辊压成膜”的干法路线：把石墨、导电炭黑和 PTFE 粘结剂像揉面一样干混，再让 PTFE 拉出纤维把粉体“缠”成整体，用自建的六辊压延机连续压出 600 毫米宽的电极。装成的软包电池能量密度更高、更耐用、低温表现也好，为无溶剂的下一代电池产线提供了落地样本。",
2: "铁-空气电池便宜、安全、材料遍地都是，是储能的好苗子，但铁在充放电时体积剧烈胀缩、离子在固体内走不动，导致性能很差。研究者的思路是在电极里搭一张“三合一网络”：掺入 LSM 和 YSZ 两种陶瓷，让离子和电子各有通畅的道路，界面还能帮忙“松开”铁氧键，刚性的骨架再给膨胀留出缓冲。结果面容量提升 14 倍、稳定运行 1000 小时。",
3: "一呼一吸之间藏着心肺、代谢甚至情绪的大量信息，但医院的肺量计只能做“一次性拍照”，无法连续记录。这篇综述描绘了下一代方案：像皮肤一样柔软贴合的可穿戴应变/湿度传感器负责“录音”，人工智能负责“听音辨病”，在家就能测肺功能、发现异常信号。文章同时指出材料设计、临床验证和数据隐私是落地的三道坎。",
4: "细胞膜由成千上万种脂质拼成，为什么配比这么讲究？研究者的巧思是把酵母和人类细胞压到深海级别的 250 个大气压下——高压会把磷脂“压扁”，等于从物理上改变了膜的形状需求。结果发现真核细胞会主动多合成“更弯”的脂质来抵消压力，而细菌不会。这说明细胞像调音师一样，主动把膜维持在一个“微微紧绷”的状态，这对膜运输等核心功能很重要。",
5: "铝遇到水本可以剧烈反应放出氢气，是潜在的“化学燃料”，但表面那层几纳米的氧化膜让它处于钝化状态。MIT 团队先把铝粒浸入液态镓铟合金“激活”，再用氨水代替纯水，结果在相当于 550 米深海的高压下依然高效产氢。这给水下航行器的浮力引擎提供了能量密度远超电池的“水下打气筒”。",
6: "做靶向药物递送经常遇到尴尬：配体在人细胞上好用，到了小鼠模型却抓不住对应的受体，动物实验做不成。这篇论文找到原因——小鼠 langerin 受体多出一块结构挡住了配体，而只要把配体的一个手性中心“翻个面”（换成 α 异头体），就能绕开障碍，同时保留对人的亲和力，实现跨物种递送。",
7: "超级电容器需要电极又导电又能存电荷。这项研究给氧化锌纳米棒同时掺入钴和铜两种金属“杂质”，把材料的带隙从 3.11 电子伏压窄到 2.15 电子伏，电荷转移更快、储能更多。实验表征与 DFT 计算相互印证：掺入的金属在费米能级附近提供了新的电子态，是带隙变窄的根源。",
8: "“水伏发电”靠水与材料表面的相互作用产生电，微生物生物膜是廉价的发电平台，可惜功率太小。这项研究给希瓦氏菌“穿上一层聚苯胺外套”：细菌与导电聚合物之间形成肖特基结，电子流出更快、内阻更低，功率密度提升一个数量级，让“用细菌发电”离实用更近。",
9: "“催化剂被光促进时，该用一种光还是几种光？”这篇理论工作用动力学蒙特卡洛模拟给出了漂亮答案：当催化循环里有两个慢步骤时，用两种波长的光分别去加速它们，总速率提升最大——好比两个人各推一辆陷住的车。研究还绘制了“光子通量-量子产率”相图，告诉实验者什么条件下发光效率最高。",
10: "“肌肉有记忆”广为人知，细胞其实也有“力学记忆”：被基质挤压、牵拉过的细胞会在表观遗传层面留下痕迹，影响之后的命运决定。这篇综述系统梳理了机械记忆如何被写入（染色质改变）、存住（细胞骨架与核架构）和擦除，并讨论如何据此设计更好的植入材料与个性化疗法。",
11: "“手性”是药物的命门——镜像分子药效可能天差地别。研究者合成了 16 种带手性中心的磺酰亚胺基氟“弹头”，测试它们对模型酶胰凝乳蛋白酶的抑制能力：最活跃的对映体比它的镜像兄弟强 29–54 倍。X 射线晶体学、质谱肽图和分子对接共同解释了差距来源：只有“对的手性”才能让弹头以合适角度够到活性位点的丝氨酸。",
12: "如何在纳米尺度精确控制多种稀土的配比？研究者的答案是把反应“关进”脂质体：这个几十纳米的“反应器”体积只有阿升级，转运体把配体运进去，Tb、Eu 离子在内部缓慢配位成发光配合物。调节两者比例就能调出从绿到红的颜色，再加上蓝色配体凑齐全 RGB。有趣的是，实时监测发现 Eu 的发光总是先亮起来。",
13: "“试纸变个色就能看出有没有感染”是诊断的理想形态：不用仪器、便宜、快速。这篇观点文章以作者团队 15 美分一次的新冠 COLOR 检测为案例，系统梳理比色生物传感的材料工具箱（金纳米颗粒、适配体、纳米酶等），并直言四大短板：稳定性、选择性、重现性和规模化生产——给出从实验室走向家庭的路线图。",
14: "肽核酸（PNA）是能钻进双链 DNA 的“分子夹”，有望直接关掉癌基因，但一直没人在真实染色体上验证。这项研究比较两种设计后发现：带“G-clamp”碱基修饰的 PNA 能在癌细胞里序列特异性地夹住 c-Myc 基因、压低蛋白表达、减缓增殖；先用 HDAC 抑制剂把染色质“放松”，低剂量 PNA 也能起效。",
15: "大脑里的星形胶质细胞像“管家”，调控着神经元的兴奋程度。这项研究做了一个巧妙的对照：碳纳米管和硅酸盐纳米管外形几乎一样，一个导电一个不导电。结果只有导电版本能让星形胶质细胞钙活动增强、谷氨酸摄取提升，进而降低神经元放电——说明起作用的是“导电性”而非纳米形状。",
16: "治抑郁的电极要么开颅、要么刺激不精准。这项研究把超薄透明电极做进隐形眼镜，利用“时间干涉”原理：两路高频电流在视网膜交叉处合成低频包络，只刺激深部脑区而表皮无感。抑郁模型小鼠戴上它刺激几周，行为、脑区间同步和分子标志物都明显恢复，机器学习分类证实疗效接近药物氟西汀。",
17: "风电、光伏让电网变成一个巨大的软件系统，攻击面随之暴增。文章梳理了 20 年来 20 起标志性事件（从 Stuxnet 到乌克兰电网攻击），并用仿真展示两类狠角色：虚假数据注入让电压越限，针对构网型变流器的同步攻击能让低惯量电网频率振荡。作者提出零信任架构等六项优先事项。",
18: "化学家看分子靠“图”，机器读分子靠“文本”（SMILES），两者长期割裂。ChemMLLM 是第一个同时打通三种模态的化学大模型：用图像分词器把分子结构图变成 token，让语言模型统一学习“看图说话、看图猜性质、按性质画分子”。五类任务上全面超越 GPT-4o 等通用模型，还通过注意力热图证明它真的在看官能团。",
19: "环氧树脂又强又耐用，坏处是“寿终”后难以回收。这项研究设计了一类室温就能固化的动态硼酸酯交联剂：做成的“类玻璃体”塑料加热可以像热塑性材料一样重压成型，泡在乙醇水里又能解聚回收，反复四次性能不掉。制成的碳纤维复合材料性能媲美商业环氧，还能把树脂和纤维双双完整回收。",
20: "电器插头、开关为什么会坏在接触点边缘？研究者建立了微凸体尺度的电-热-力全耦合仿真：电流流过微小接触面时在边缘“挤成一股细流”，电流密度趋于发散，局部焦耳热让边缘软化，塑性变形把材料挤成“堆堆”，裂纹就从这里萌生。载流压痕实验直接看到了这些边缘裂纹与堆积，证实了理论预测。",
21: "酰胺键是蛋白质和半数药物的核心连接，但传统成酰胺方法依赖活化试剂、产生大量废物。这项研究开发了一类“烯醇酯”试剂：本身像老酒一样能存放一年以上，在水里 37 度、不加催化剂就能把羧酸和胺连起来，副产物二酮过滤回收后还能再做成新试剂，实现“循环”合成，连蛋白质赖氨酸修饰都不在话下。",
22: "AI 生成分子的能力越来越强，但它到底学会了化学，还是只是“背答案”？DiffSHAPer 把博弈论里的 Shapley 值引入分子扩散模型：把每个原子当成“玩家”，计算它对生成结果的贡献。审计结论出人意料——连接子生成模型并没有学到可泛化的化学规则，主要靠锚原子附近的距离约束在起作用。",
23: "纳米药物进入肿瘤细胞后有个隐形敌人：细胞每分裂一次，吃进去的颗粒就被两个子细胞“对半分”，浓度不断被稀释。研究者的策略是“先踩刹车再加油”：用放疗或 CDK 抑制剂把细胞按在分裂前，再给纳米药，颗粒滞留时间显著延长。小鼠实验里，多柔比星脂质体加一次照射，药物半衰期从约 140 小时延长到 248 小时。",
24: "脑出血后最难恢复的是白质损伤——髓鞘坏了，信号传不动。这项研究把两种疗法装进一个载体：氧化铈纳米粒负责清除自由基、调节炎症，上面再负载多发性硬化药物芬戈莫德促进髓鞘修复。纳米粒经 STAT3 通路推动少突胶质前体细胞成熟，小鼠的运动和记忆功能明显改善，为“一石二鸟”的神经纳米药提供了范例。",
25: "“这个分子有多想吃电子对？”对带两个正电荷的“双头”Lewis 酸，一直缺少统一标尺。这项计算研究提出用氧离子亲和能（OIA）作为新指标，系统计算了 164 个双阳离子，发现 OIA 与常用的氟离子亲和能（FIA）强相关、可直接换算，还能预测哪些双鏻盐会从三氟甲磺酸根那里“抢氧”发生脱氧反应。",
26: "卫星太阳能板的保护玻璃被空间电子辐照后会变黑、发电变差，传统解法是加铈，又贵又重。这项研究发现：不加铈的商用硼硅玻璃在电子辐照致黑后，只要晒模拟太阳光就能“自我漂白”，缺陷湮灭速度远快于生成速度——在轨“辐照与光照并存”的条件下可长期保持透亮，是低成本盖片方案。",
27: "“零衰减电池”可能吗？关键在于补上电池变旧的根源——活性锂的持续损失。研究者在 PyBaMM 物理模型中给电池加了一种“补锂剂”（LFO），系统试了五种锂释放策略和多个含量：结论是要实现零衰减，锂必须慢慢放、时机要卡准、含量要最优（约 4%），还得给负极留出冗余容量，否则按下葫芦浮起瓢。",
28: "风电便宜但“看天吃饭”，出力低谷时缺口怎么办？研究者把美国 583 个风电场的月度发电数据与地下 5 公里深处增强型地热的资源地图做空间配对：56% 的陆上风电场脚下就埋着足以补齐缺口的地热，且其中 104 个场址靠近城市和电网。风电+地热同址共建，有望把“靠天吃饭”的风变成稳定电源。",
29: "让“分子夹”夹住双链 RNA 的诀窍藏在一个原子里：在识别特殊碱基对的修饰旁边再引入一个 2-硫代尿嘧啶，三链体稳定性大增——主要不是“抓得更紧”，而是“松手更慢”（解离速率下降）。这一招能增强核糖体移码效率，还能推广到 miRNA 前体和甲流病毒 RNA 的靶向，是设计 RNA 结合探针的通用思路。",
30: "水泥是人类用量最大的人造材料，却基本是“死”的。这篇观点文章描绘了把它变成微生物“家园”的路线图：能自愈裂缝的细菌孢子、固碳的蓝细菌、发电的地杆菌、防钢筋锈蚀的菌群、感知环境的生物传感——并指出高碱度、孔隙结构和无法原位观测是三大瓶颈，未来要靠基因工程与 AI 设计“活的建材”。",
31: "器官冷冻保存依赖防冻剂阻止冰晶形成，但为什么有的分子好用、有的不行，一直没有分子层面的解释。研究者把低温分子动力学模拟与实验测得的“玻璃化最低浓度”关联起来，发现关键在于防冻剂破坏水分子四面体氢键网格的能力——破坏得越彻底越容易形成玻璃态。从此设计新防冻剂可以“按图索骥”而非大海捞针。",
32: "直接从空气里抓 CO2，最怕吸附剂又贵又娇气。氧化石墨烯是理想候选：比表面积大、介孔多、表面富含含氧基团。靠环境湿度变化就能可逆地吸放 CO2，吸附容量 1.354 mmol/g，循环 1000 次几乎不衰减。研究还演示了一个农用场景：把吸附膜挂在桑田里，释放的 CO2 就地当“气肥”，桑叶都长得更重了。",
33: "纳塑料小到无法过滤、少到难以检测，是水处理的新难题。研究者的灵感来自水母黏液：合成一种两亲性瓶刷聚合物，常温下是溶液，加热就变成凝胶并“收缩挤水”，把疏水的纳塑料富集进凝胶网络；冷却后又释放出来供检测，去除率 68%–100%。这套“抓了再放”的可逆捕集器还能靠带电基团选择性抓染料。",
34: "碳捕集最烧钱的一步是解吸——把 CO2 从吸收剂里“烤”出来需要 120 度高温。研究者的巧思是把化工里的“萃取”搬过来：设计一种吸收时溶为一体、解吸时自动分两层的胺体系，CO2 产物被另一组分“拉”出来，解吸温度直降 30–40 度，工厂余热即可驱动，再生能耗比传统 MEA 低 62%。",
35: "3D 闪存越做越密，数据却越来越容易“串门”——电荷横向迁移到相邻单元导致数据丢失。研究者的方案是给电荷建“独栋牢房”：用原子层沉积交替生长三层氧化铪纳米晶，每个纳米晶都被氧化铝绝缘层包围，横向纵向都跑不出去。存储窗口 7 伏、3-bit 数据保持率 86%，工艺还兼容高深宽比的 3D 架构。",
36: "现代药物常需要同时作用于多个靶点，但传统构效关系一次只看一个靶点。这篇综述盘点了一套开源化学信息学“工具箱”——12 个能从多靶点大数据集中挖掘“结构-多重活性关系”的平台，配合可视化与网络分析帮研究者在效力、安全与耐药之间找平衡，最后用抗菌肽设计实战演示了完整工作流。",
37: "抗病毒涂层靠氧化灭活病毒，但它到底把病毒怎么了？研究者用原子力显微镜对单个病毒做“体检”：接触 N-卤胺涂层后，烟草花叶病毒的形状没变、RNA 完好，蛋白外壳却明显变软、更容易被压破——像抽掉了帐篷的钢架。计算模拟还定位了最容易被氧化的衣壳残基，为设计更强的抗病毒表面指明靶点。",
38: "类器官是比细胞培养更真实的“迷你器官”，但依赖昂贵的基质胶。这项研究用便宜的羧甲基纤维素水凝胶，装上细胞抓手的 RGD 肽、调节软硬，就能加速多种类器官形成，其中的整合素/FAK/YAP 通路是关键。团队还建立了“走路摩擦的机械应力驱动肢端黑色素瘤”模型，为低成本的药物筛选提供了平台。",
39: "房车是移动的家，但密闭车厢里的空气 quality 如何？研究者在苏州对一辆房车做了 32 天实境观测：甲醛是头号污染物（均值超标 WHO 限值），且释放呈“长尾”——材料表面温度越高，释放越猛。团队用 SHAP 可解释性加 LSTM 深度学习模型预测浓度变化并评估健康风险，为移动居住空间的环境设计提供了依据。",
40: "车轮转动能发电吗？研究者把碳纳米管纱线捻成“小弹簧”，装进跷跷板结构的收集器：车轮一转，两根纱线交替拉伸、回弹，其电化学电容随之变化而输出电流。设备在 1500 转每分、20–60 度、40%–80% 湿度和颠簸冲击下都稳定发电，经升压电路可充到 4.2 伏——足够给胎压监测这类车载传感器供电，摆脱电池更换烦恼。",
41: "发酵罐里的酵母看起来整齐划一，其实每个细胞的“产量”参差不齐。研究者用高光谱受激拉曼成像给单个酵母细胞无损“称重”瓦伦烯含量，再配合 Lasso 解混算法把每种分子的浓度画成图：不仅看清了细胞间的巨大差异，还发现细胞优先“吃”乙酸而不是葡萄糖来合成产物。据此优化碳氮比和补料策略，产量提到 125 mg/L。",
42: "蛋白质与 DNA 的结合往往一瞬即逝，怎么抓个“现行”？研究者在 G-四链体 DNA 里预埋一个呋喃“暗扣”，再让光敏剂就地产生单线态氧——光照一下，呋喃被激活，与紧挨着的结合肽上的赖氨酸形成稳定共价锁。整个反应由光时空可控，且锁上之后 G4 结构安然无恙，为解析瞬时的蛋白-DNA 相互作用按下“快门”。",
43: "钠电池负极选材是笔“糊涂账”——文献里上百种二维材料各说各话。这项研究用高通量 DFT 一口气算了 63 种层状过渡金属硫属化物的嵌钠电压，画出一张“电压地图”：VIB 族（钼、钨系）电压最低但有相变风险，而 MoSe2、WSe2、MoTe2、WTe2 兼具合适电压与相稳定性，是最值得实验验证的候选。",
44: "把催化剂“种”在多孔材料里，微环境决定活性。研究者在聚离子液体框架中对比两种形状的交联剂：笼状 POSS 与平面星形分子。结果笼状结构不仅促使钌聚集成大颗粒，还带来局部氯富集，造就“缺电子”的钌位点——恰好是氢甲酰化反应最喜欢的状态，转化数高达 1038。“交联剂形状调控微环境”由此成为负载催化剂的新设计维度。",
45: "太阳能直接分解水制绿氢，被视为终极清洁能源方案之一，但效率始终上不去。这期 Voices 邀集全球专家把脉：瓶颈集中在非氧化物光催化剂（氮化物、硫氧化物）的阴离子缺陷——它们让载流子活不长；解决之道包括把缺陷浓度降到氧化物体系水平、发展工作态（operando）表征、从间歇式走向连续化生产，以及材料、器件、工程多学科协同。",
46: "AI 芯片越来越烫，传统液冷冷板的散热片形状受限于加工工艺。研究者让拓扑优化算法“放飞”设计出最优流道，再用“电化学增材制造”直接把纯铜冷板打印出来——亚 100 微米的精细结构成为可能。实测热阻降低 32%、同散热下压降降 68%，数据中心冷却能耗占比可压到 1.1%。关键是设计与制造终于“门当户对”。",
47: "细胞里有一类没有膜的“细胞器”，像油和水分离一样靠液-液相分离形成。悖论在于：内部蛋白浓度高得惊人，却还要保持液体般的流动。全原子模拟揭示了答案：无序蛋白之间隔着大量“活力水分子”，氢键快速交换，让接触既粘又滑——既维持流动性，又防止蛋白骨架抱死成有序结构（病理纤维的前兆）。",
48: "甲醇变丙烯是化工的重要路线，但催化剂容易积碳“噎死”。一维大孔的 ZSM-12 沸石是个异类：丙烯选择性 40%–55%，累计吞吐量高达 480 克甲醇每克催化剂。固体核磁、红外、紫外可见光谱与积碳萃取共同揪出了环戊烯基、甲基苯鎓等关键中间体；中子衍射则直接“看到”焦炭堵孔与晶格膨胀的全过程，解释了它为什么这么能扛。",
}

# ---------------- Markdown 解析 ----------------
def parse_batch(paths):
    papers = {}
    for path in paths:
        with open(path, encoding="utf-8") as f:
            text = f.read()
        cur = None
        for line in text.splitlines():
            m = re.match(r"^##\s*(\d+)\.\s*(.+)$", line)
            if m:
                cur = {"num": int(m.group(1)), "title": clean(m.group(2)),
                       "domain": "", "overview": "", "methods": [], "vocab": []}
                papers[cur["num"]] = cur
                continue
            if cur is None:
                continue
            m = re.match(r"^-\s*\*\*领域\*\*\s*[:：]\s*(.+)$", line)
            if m:
                cur["domain"] = clean(m.group(1)); continue
            m = re.match(r"^-\s*\*\*一句话概述\*\*\s*[:：]\s*(.+)$", line)
            if m:
                cur["overview"] = clean(m.group(1)); continue
            if re.match(r"^-\s*\*\*方法论总结\*\*", line):
                cur["_in"] = "m"; continue
            if re.match(r"^-\s*\*\*英语学习重点词汇\*\*", line):
                cur["_in"] = "v"; continue
            m = re.match(r"^\s*\d+\.\s+(.+)$", line)
            if m and cur.get("_in") == "m" and not line.startswith("|"):
                cur["methods"].append(clean(m.group(1))); continue
            if line.startswith("|"):
                cells = [c.strip() for c in line.strip("|").split("|")]
                if any(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    continue
                if cells and (cells[0] in ("词汇", "词汇 ", "词汇（学术）") or cells[0].startswith("词汇")):
                    continue
                if len(cells) >= 3 and cells[0]:
                    cur["vocab"].append([clean(cells[0], quote=True), clean(cells[1], quote=True),
                                         clean(cells[2], quote=True)])
                continue
    return [papers[k] for k in sorted(papers)]

def clean(s, quote=False):
    s = s.replace("**", "").strip()
    if quote:
        s = s.strip().strip('"').strip().strip('"').strip()
    return s

# ---------------- docx 底层工具 ----------------
def el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), str(v))
    return e

def set_run_font(run, cn="SimSun", en="Times New Roman", size=12, bold=False,
                 color="000000", italic=False):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = el("w:rFonts"); rpr.insert(0, rf)
    rf.set(qn("w:eastAsia"), cn)

def add_text(p, text, **kw):
    r = p.add_run(text)
    set_run_font(r, **kw)
    return r

def para_shading(p, fill):
    p._p.get_or_add_pPr().append(el("w:shd", val="clear", color="auto", fill=fill))

def para_border(p, side, sz=6, color=ACCENT, space=4):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = el("w:pBdr"); pPr.append(pBdr)
    pBdr.append(el("w:" + side, val="single", sz=str(sz), space=str(space), color=color))

def tbl_borders_none(table):
    tblPr = table._tbl.tblPr
    b = el("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b.append(el("w:" + side, val="none", sz="0", space="0", color="auto"))
    tblPr.append(b)

def tbl_borders_h(table, color=T_INNER, sz=4):
    tblPr = table._tbl.tblPr
    b = el("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideV"):
        b.append(el("w:" + side, val="none", sz="0", space="0", color="auto"))
    b.append(el("w:insideH", val="single", sz=str(sz), space="0", color=color))
    tblPr.append(b)

def tbl_width_pct(table):
    tblPr = table._tbl.tblPr
    tblPr.append(el("w:tblW", w="5000", type="pct"))
    tblPr.append(el("w:tblLayout", type="fixed"))

def tbl_cell_margins(table, top=40, left=100, bottom=40, right=100):
    tblPr = table._tbl.tblPr
    m = el("w:tblCellMar")
    for tag, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        m.append(el("w:" + tag, w=str(val), type="dxa"))
    tblPr.append(m)

def cell_shade(cell, fill):
    cell._tc.get_or_add_tcPr().append(el("w:shd", val="clear", color="auto", fill=fill))

def cell_width_pct(cell, pct):  # pct: 0-100
    cell._tc.get_or_add_tcPr().append(el("w:tcW", w=str(int(pct * 50)), type="pct"))

def row_cant_split(row, header=False):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(el("w:cantSplit"))
    if header:
        trPr.append(el("w:tblHeader"))

def cell_p(cell, first=True):
    return cell.paragraphs[0] if first and cell.paragraphs else cell.add_paragraph()

def page_field(p, instr):
    r = p.add_run(); r._r.append(el("w:fldChar", fldCharType="begin"))
    r2 = p.add_run(); it = el("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r2._r.append(it)
    r3 = p.add_run(); r3._r.append(el("w:fldChar", fldCharType="separate"))
    r4 = p.add_run("1")
    r5 = p.add_run(); r5._r.append(el("w:fldChar", fldCharType="end"))
    for rr in (r, r2, r3, r4, r5):
        set_run_font(rr, cn="SimSun", size=9, color="595959")

def set_pgnum(section, fmt=None, start=None):
    sectPr = section._sectPr
    e = sectPr.find(qn("w:pgNumType"))
    if e is None:
        e = el("w:pgNumType"); sectPr.append(e)
    if fmt:
        e.set(qn("w:fmt"), fmt)
    if start is not None:
        e.set(qn("w:start"), str(start))

def tiny_para(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    r = p.add_run(""); set_run_font(r, size=2)
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(2)

# ---------------- 封面（R1 配方，python-docx 实现） ----------------
def split_title_lines(title, units_per_line):
    def width(s):
        return sum(1.0 if ord(c) > 0x2E80 else 0.55 for c in s)
    if width(title) <= units_per_line:
        return [title]
    break_after = set("，。、；：！？的与和及之在于为-_—–·/ \t")
    lines, remaining = [], title
    while width(remaining) > units_per_line:
        acc, brk = 0.0, -1
        for i, ch in enumerate(remaining):
            if acc >= units_per_line * 0.6 and brk < 0 and ch in break_after:
                brk = i + 1
            if acc >= units_per_line:
                if brk < 0:
                    brk = i
                break
            acc += 1.0 if ord(ch) > 0x2E80 else 0.55
        if brk <= 0:
            brk = max(1, int(units_per_line))
        lines.append(remaining[:brk].strip())
        remaining = remaining[brk:].strip()
    if remaining:
        lines.append(remaining)
    if len(lines) > 1 and len(lines[-1]) <= 2:
        last = lines.pop(); lines[-1] += last
    return lines

def calc_cover_spacing(title_lines, title_pt, has_sub, has_label, n_meta):
    safety = 1200
    usable = 16838 - safety
    content = (title_lines * (title_pt * 23 + 200)
               + (12 * 23 + 600 if has_sub else 0)
               + (9 * 23 + 600 if has_label else 0)
               + n_meta * (10 * 23 + 100) + 400 + 3 * 300)
    remaining = max(usable - content, 400)
    foot_min = 800
    raw_top = remaining * 45 // 100
    raw_bottom = remaining * 45 // 100
    bottom = max(raw_bottom, foot_min)
    top = max(raw_top - max(0, foot_min - raw_bottom), 400)
    return top, bottom

def build_cover(doc, batch_cn, batch_no, start, end, n):
    title = "Cell Press 论文精读报告"
    label = "P A P E R   R E A D I N G   N O T E S"
    subtitle = f"批次{batch_cn} · 论文 {start:02d}–{end:02d} · 方法论总结与英语学习词汇"
    metas = [f"来源期刊：Cell Reports Physical Science（2026 年 5 月刊）",
             f"收录论文：{n} 篇（编号 {start:02d}–{end:02d}）",
             "内容：领域概述 · 通俗解读 · 方法论总结 · 重点词汇",
             "生成日期：2026-09-01"]

    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl_width_pct(tbl); tbl_borders_none(tbl); tbl_cell_margins(tbl, 0, 1200, 0, 800)
    row = tbl.rows[0]
    trPr = row._tr.get_or_add_trPr()
    trPr.append(el("w:trHeight", val="16700", hRule="exact"))
    cell = row.cells[0]
    cell_shade(cell, BG)

    title_pt = 36
    lines = split_title_lines(title, (11906 - 1200 - 800 - 300) / (title_pt * 20))
    top_sp, bottom_sp = calc_cover_spacing(len(lines), title_pt, True, True, len(metas))

    p = cell_p(cell); p.paragraph_format.space_before = Twips(top_sp); p.paragraph_format.space_after = Pt(0)
    p2 = cell_p(cell, first=False)
    p2.paragraph_format.space_after = Pt(18)
    para_border(p2, "bottom", sz=6, color=ACCENT, space=6)
    add_text(p2, label, cn="SimHei", en="Calibri", size=9, color=ACCENT)

    for ln in lines:
        tp = cell_p(cell, first=False)
        tp.paragraph_format.space_after = Pt(4)
        tp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        tp.paragraph_format.line_spacing = Pt(math.ceil(title_pt * 23 / 20))
        add_text(tp, ln, cn="SimHei", en="SimHei", size=title_pt, bold=True, color=C_TITLE)

    sp = cell_p(cell, first=False); sp.paragraph_format.space_before = Pt(14); sp.paragraph_format.space_after = Pt(26)
    add_text(sp, subtitle, cn="SimSun", size=12, color=C_SUB)

    for m in metas:
        mp = cell_p(cell, first=False)
        mp.paragraph_format.space_after = Pt(7)
        para_border(mp, "left", sz=8, color=ACCENT, space=6)
        add_text(mp, m, cn="SimSun", size=10, color=C_META)

    fp = cell_p(cell, first=False); fp.paragraph_format.space_before = Twips(bottom_sp)
    fp.paragraph_format.space_after = Pt(0)
    para_border(fp, "top", sz=6, color=ACCENT, space=8)
    add_text(fp, "Cell-Press-Paper-Analysis-ChemMLLM  ·  2026-09", cn="SimHei", en="Calibri", size=9, color=C_FOOT)
    return tbl

# ---------------- 正文构件 ----------------
def h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.3
    add_text(p, text, cn="SimHei", en="Times New Roman", size=16, bold=True, color=H1_COLOR)
    return p

def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    add_text(p, text, cn="SimHei", en="Times New Roman", size=14, bold=True, color="000000")
    return p

def body_para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Twips(480)
    add_text(p, text)
    return p

def info_table(doc, domain, overview):
    t = doc.add_table(rows=2, cols=2)
    t.autofit = False
    tbl_width_pct(t); tbl_borders_h(t); tbl_cell_margins(t, 60, 120, 60, 120)
    widths = (14, 86)
    data = [("领域", domain), ("一句话概述", overview)]
    for i, (label, val) in enumerate(data):
        row = t.rows[i]; row_cant_split(row)
        c0, c1 = row.cells
        cell_width_pct(c0, widths[0]); cell_width_pct(c1, widths[1])
        cell_shade(c0, T_SURF)
        p0 = cell_p(c0); add_text(p0, label, cn="SimHei", size=10.5, bold=True, color=H1_COLOR)
        p1 = cell_p(c1); add_text(p1, val, size=10.5)
        for p in (p0, p1):
            p.paragraph_format.line_spacing = 1.3; p.paragraph_format.space_after = Pt(0)
    return t

def callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    para_border(p, "left", sz=16, color=ACCENT, space=4)
    para_shading(p, T_SURF)
    add_text(p, label + "　", cn="SimHei", size=11, bold=True, color=H1_COLOR)
    add_text(p, text, size=11)
    return p

def numbered(doc, items):
    for i, text in enumerate(items, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.line_spacing = 1.3; pf.space_after = Pt(4)
        pf.left_indent = Twips(480); pf.first_line_indent = Twips(-480)
        add_text(p, f"{i}. ", cn="SimHei", size=12, bold=True, color=H1_COLOR)
        add_text(p, text)

def vocab_table(doc, rows):
    t = doc.add_table(rows=1 + len(rows), cols=3)
    t.autofit = False
    tbl_width_pct(t); tbl_borders_h(t); tbl_cell_margins(t, 50, 100, 50, 100)
    widths = (20, 28, 52)
    heads = ("词汇", "释义", "论文中的语境")
    for j, htxt in enumerate(heads):
        c = t.rows[0].cells[j]
        cell_width_pct(c, widths[j]); cell_shade(c, T_HEADBG)
        p = cell_p(c); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.3
        add_text(p, htxt, cn="SimHei", size=10.5, bold=True, color="FFFFFF")
    row_cant_split(t.rows[0], header=True)
    for i, (w, m, ctx) in enumerate(rows, 1):
        row_cant_split(t.rows[i])
        for j, val in enumerate((w, m, ctx)):
            c = t.rows[i].cells[j]
            cell_width_pct(c, widths[j])
            p = cell_p(c); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.3
            add_text(p, val, size=10.5, bold=(j == 0))
    return t

def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(16)
    add_text(p, "目　录", cn="SimHei", size=16, bold=True)
    fld = doc.add_paragraph()
    r = fld.add_run(); r._r.append(el("w:fldChar", fldCharType="begin", dirty="true"))
    r2 = fld.add_run(); it = el("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = ' TOC \\o "1-1" \\h \\z \\u '; r2._r.append(it)
    r3 = fld.add_run(); r3._r.append(el("w:fldChar", fldCharType="separate"))
    r4 = fld.add_run("（目录页码将在打开文档后自动生成）"); set_run_font(r4, size=10.5, color="808080")
    r5 = fld.add_run(); r5._r.append(el("w:fldChar", fldCharType="end"))
    hint = doc.add_paragraph()
    hint.paragraph_format.space_before = Pt(10)
    add_text(hint, "提示：在 Word / WPS 中打开本文档后，右键点击目录区域并选择“更新域”，即可显示完整目录与页码。",
             size=9, italic=True, color="808080")

# ---------------- 组装 ----------------
BATCHES = [
    ("1", "一", "01", "10", ["batch1.md", "batch1b.md"]),
    ("2", "二", "11", "20", ["batch2.md"]),
    ("3", "三", "21", "30", ["batch3.md"]),
    ("4", "四", "31", "40", ["batch4a.md", "batch4b.md"]),
    ("5", "五", "41", "48", ["batch5a.md", "batch5b.md"]),
]

def build_one(batch_no, batch_cn, start, end, files):
    papers = parse_batch([os.path.join(SUM, f) for f in files])
    n = len(papers)
    fname = f"批次{batch_no}_论文{start}-{end}_方法论与英语学习词汇.docx"
    path = os.path.join(OUT, fname)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    npf = normal.paragraph_format
    npf.line_spacing = 1.3; npf.space_after = Pt(6)

    # 封面节
    sec = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, attr, 0)
    build_cover(doc, batch_cn, batch_no, int(start), int(end), n)

    # 目录节（罗马页码）
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    for attr, v in (("top_margin", 1440), ("bottom_margin", 1440), ("left_margin", 1701), ("right_margin", 1417)):
        setattr(sec2, attr, Twips(v))
    tiny_para(doc.paragraphs[-1])
    add_toc(doc)
    set_pgnum(sec2, fmt="upperRoman", start=1)
    sec2.footer.is_linked_to_previous = False
    fpar = sec2.footer.paragraphs[0]; fpar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_field(fpar, r"PAGE \* ROMAN \* MERGEFORMAT")

    # 正文节（阿拉伯页码，从 1 起）
    sec3 = doc.add_section(WD_SECTION.NEW_PAGE)
    for attr, v in (("top_margin", 1440), ("bottom_margin", 1440), ("left_margin", 1701), ("right_margin", 1417)):
        setattr(sec3, attr, Twips(v))
    tiny_para(doc.paragraphs[-1])
    set_pgnum(sec3, fmt="decimal", start=1)
    sec3.header.is_linked_to_previous = False
    hd = sec3.header.paragraphs[0]; hd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(hd, f"Cell Press 论文精读报告 · 批次{batch_cn}", cn="SimHei", size=9, color="808080")
    para_border(hd, "bottom", sz=4, color="BFBFBF", space=2)
    sec3.footer.is_linked_to_previous = False
    fpar = sec3.footer.paragraphs[0]; fpar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_field(fpar, r"PAGE \* arabic \* MERGEFORMAT")

    # 正文
    intro = ("本批次收录论文 " + start + "–" + end + "，共 " + str(n) + " 篇，覆盖领域包括："
             + "、".join(sorted({p["domain"].split("/")[0].strip() for p in papers})) + "。"
             + "每篇论文依次呈现领域概述、通俗解读、方法论总结与英语学习重点词汇，"
               "适合先读“通俗解读”建立整体印象，再对照“方法论总结”深入技术细节，"
               "最后借助词汇表积累学术英语表达。")
    body_para(doc, intro)

    for k, pa in enumerate(papers):
        num = pa["num"]
        title_p = h1(doc, f"论文 {num:02d}　{pa['title']}")
        if k > 0:
            title_p.paragraph_format.page_break_before = True
        info_table(doc, pa["domain"], pa["overview"])
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
        callout(doc, "通俗解读", PLAIN.get(num, ""))
        h2(doc, "方法论总结")
        numbered(doc, pa["methods"])
        h2(doc, "英语学习重点词汇")
        vocab_table(doc, pa["vocab"])

    doc.save(path)
    return path, n

if __name__ == "__main__":
    total = 0
    for b in BATCHES:
        path, n = build_one(*b)
        total += n
        print(f"OK  {os.path.basename(path)}  ({n} papers)")
    print("total papers:", total)
